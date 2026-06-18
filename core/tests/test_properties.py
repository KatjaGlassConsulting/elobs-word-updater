from __future__ import annotations
import io
from unittest.mock import patch
from lxml import etree
from docx import Document
from elobs_word_updater.document.properties import (
    get_custom_property,
    set_custom_property,
    update_study_properties,
    CUSTOM_PROPS_NS,
    VT_NS,
)


def _roundtrip(doc: Document) -> Document:
    """Save the doc to an in-memory buffer and reopen it (real persistence path)."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


# --- Real persistence tests (no mocks): document has no custom-properties part ---

def test_set_custom_property_creates_part_and_persists():
    """A property set on a doc with no custom.xml part is created and survives save/reload."""
    doc = Document()  # default template has no docProps/custom.xml
    set_custom_property(doc, "StudyUID", "Study_000001")
    reloaded = _roundtrip(doc)
    assert get_custom_property(reloaded, "StudyUID") == "Study_000001"


def test_update_study_properties_persists_all_five():
    """update_study_properties creates and persists all five properties on a fresh document."""
    doc = Document()
    update_study_properties(doc, "Study_000001", "CDISC-001", "3", "RELEASED", "2026-03-07T11:00:00Z")
    reloaded = _roundtrip(doc)
    assert get_custom_property(reloaded, "StudyUID") == "Study_000001"
    assert get_custom_property(reloaded, "StudyId") == "CDISC-001"
    assert get_custom_property(reloaded, "StudyVersion") == "3"
    assert get_custom_property(reloaded, "StudyVersionStatus") == "RELEASED"
    assert get_custom_property(reloaded, "OSBSyncedAt") == "2026-03-07T11:00:00Z"


def _make_custom_props_root() -> etree._Element:
    return etree.Element(f"{{{CUSTOM_PROPS_NS}}}Properties")


def _add_prop(root: etree._Element, name: str, value: str) -> None:
    prop = etree.SubElement(root, f"{{{CUSTOM_PROPS_NS}}}property")
    prop.set("fmtid", "{D5CDD505-2E9C-101B-9397-08002B2CF9AE}")
    prop.set("pid", "2")
    prop.set("name", name)
    vt = etree.SubElement(prop, f"{{{VT_NS}}}lpwstr")
    vt.text = value


# --- Tests when no custom-properties part exists ---

def test_get_custom_property_no_part():
    """Returns None when the document has no custom properties part."""
    doc = Document()
    assert get_custom_property(doc, "StudyUID") is None


# --- Tests for get_custom_property ---

def test_get_custom_property_found():
    """Returns the value of a property that exists in the custom properties."""
    root = _make_custom_props_root()
    _add_prop(root, "StudyUID", "Study_000001")
    doc = Document()
    with patch("elobs_word_updater.document.properties._get_custom_props_element", return_value=root):
        assert get_custom_property(doc, "StudyUID") == "Study_000001"


def test_get_custom_property_not_found():
    """Returns None for a property name that is not in the custom properties."""
    root = _make_custom_props_root()
    _add_prop(root, "StudyUID", "Study_000001")
    doc = Document()
    with patch("elobs_word_updater.document.properties._get_custom_props_element", return_value=root):
        assert get_custom_property(doc, "NonExistent") is None


# --- Tests for set_custom_property ---

def test_set_custom_property_updates_existing():
    """Updates the value of an existing custom property in place."""
    root = _make_custom_props_root()
    _add_prop(root, "StudyUID", "old_value")
    doc = Document()
    with patch("elobs_word_updater.document.properties._get_custom_props_element", return_value=root), \
         patch("elobs_word_updater.document.properties._set_custom_props_element") as mock_set:
        set_custom_property(doc, "StudyUID", "Study_000001")
        mock_set.assert_called_once()
    vt = root.find(f".//{{{VT_NS}}}lpwstr")
    assert vt.text == "Study_000001"


def test_set_custom_property_adds_new_property():
    """Adds a new property element when the name does not yet exist."""
    root = _make_custom_props_root()
    doc = Document()
    with patch("elobs_word_updater.document.properties._get_custom_props_element", return_value=root), \
         patch("elobs_word_updater.document.properties._set_custom_props_element"):
        set_custom_property(doc, "StudyUID", "Study_000001")
    props = root.findall(f"{{{CUSTOM_PROPS_NS}}}property")
    assert len(props) == 1
    assert props[0].get("name") == "StudyUID"
    assert props[0].find(f"{{{VT_NS}}}lpwstr").text == "Study_000001"


def test_set_custom_property_pid_increments():
    """New property gets pid = max(existing pids) + 1."""
    root = _make_custom_props_root()
    _add_prop(root, "Existing", "val")
    root.find(f"{{{CUSTOM_PROPS_NS}}}property").set("pid", "5")
    doc = Document()
    with patch("elobs_word_updater.document.properties._get_custom_props_element", return_value=root), \
         patch("elobs_word_updater.document.properties._set_custom_props_element"):
        set_custom_property(doc, "NewProp", "newval")
    props = root.findall(f"{{{CUSTOM_PROPS_NS}}}property")
    new_prop = next(p for p in props if p.get("name") == "NewProp")
    assert int(new_prop.get("pid")) == 6


# --- Tests for update_study_properties ---

def test_update_study_properties_sets_all_five():
    """Sets all five expected properties: StudyUID, StudyId, StudyVersion, StudyVersionStatus, OSBSyncedAt."""
    root = _make_custom_props_root()
    doc = Document()
    with patch("elobs_word_updater.document.properties._get_custom_props_element", return_value=root), \
         patch("elobs_word_updater.document.properties._set_custom_props_element"):
        update_study_properties(doc, "Study_000001", "CDISC-001", "2", "FINAL", "2026-03-07T11:00:00Z")
    names = {p.get("name") for p in root.findall(f"{{{CUSTOM_PROPS_NS}}}property")}
    assert names == {"StudyUID", "StudyId", "StudyVersion", "StudyVersionStatus", "OSBSyncedAt"}


def test_update_study_properties_version_none_uses_latest():
    """When version is None, StudyVersion is set to 'LATEST' and StudyVersionStatus to 'DRAFT'."""
    root = _make_custom_props_root()
    doc = Document()
    with patch("elobs_word_updater.document.properties._get_custom_props_element", return_value=root), \
         patch("elobs_word_updater.document.properties._set_custom_props_element"):
        update_study_properties(doc, "Study_000001", "CDISC-001", None, None, "2026-03-07T11:00:00Z")
    props = {
        p.get("name"): p.find(f"{{{VT_NS}}}lpwstr").text
        for p in root.findall(f"{{{CUSTOM_PROPS_NS}}}property")
    }
    assert props["StudyVersion"] == "LATEST"
    assert props["StudyVersionStatus"] == "DRAFT"


def test_update_study_properties_values_correct():
    """All five property values are written correctly for a versioned study."""
    root = _make_custom_props_root()
    doc = Document()
    with patch("elobs_word_updater.document.properties._get_custom_props_element", return_value=root), \
         patch("elobs_word_updater.document.properties._set_custom_props_element"):
        update_study_properties(doc, "Study_000001", "CDISC-001", "3", "RELEASED", "2026-03-07T11:00:00Z")
    props = {
        p.get("name"): p.find(f"{{{VT_NS}}}lpwstr").text
        for p in root.findall(f"{{{CUSTOM_PROPS_NS}}}property")
    }
    assert props["StudyUID"] == "Study_000001"
    assert props["StudyId"] == "CDISC-001"
    assert props["StudyVersion"] == "3"
    assert props["StudyVersionStatus"] == "RELEASED"
    assert props["OSBSyncedAt"] == "2026-03-07T11:00:00Z"


def test_update_study_properties_osb_synced_at_stored():
    """The OSBSyncedAt timestamp is stored exactly as provided."""
    root = _make_custom_props_root()
    doc = Document()
    with patch("elobs_word_updater.document.properties._get_custom_props_element", return_value=root), \
         patch("elobs_word_updater.document.properties._set_custom_props_element"):
        update_study_properties(doc, "Study_000001", "CDISC-001", None, None, "2026-03-07T09:15:30Z")
    props = {
        p.get("name"): p.find(f"{{{VT_NS}}}lpwstr").text
        for p in root.findall(f"{{{CUSTOM_PROPS_NS}}}property")
    }
    assert props["OSBSyncedAt"] == "2026-03-07T09:15:30Z"
