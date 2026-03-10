from __future__ import annotations
import io
from pathlib import Path
import pytest
from docx import Document
from lxml import etree
from elobs_word_updater.document.merger import insert_docx_at_content_control

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
FIXTURES = Path(__file__).parent / "fixtures"


def _make_doc_with_cc(tag: str) -> Document:
    doc = Document()
    xml = f"""
<w:sdt xmlns:w="{W_NS}">
  <w:sdtPr><w:tag w:val="{tag}"/></w:sdtPr>
  <w:sdtContent><w:p><w:r><w:t>placeholder</w:t></w:r></w:p></w:sdtContent>
</w:sdt>""".strip()
    doc.element.body.insert(0, etree.fromstring(xml))
    return doc


def _make_simple_docx(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_insert_docx_replaces_content():
    """Inserts sub-document content into the named content control and returns True."""
    main = _make_doc_with_cc("SB_ObjectivesEndpoints")
    sub_bytes = _make_simple_docx("Inserted paragraph")
    result = insert_docx_at_content_control(main, "SB_ObjectivesEndpoints", sub_bytes)
    assert result is True
    texts = [t.text for t in main.element.body.iter(f"{{{W_NS}}}t")]
    assert any("Inserted paragraph" in (t or "") for t in texts)


def test_insert_docx_missing_tag():
    """Returns False when the target content control tag does not exist."""
    main = _make_doc_with_cc("SB_ObjectivesEndpoints")
    sub_bytes = _make_simple_docx("Some text")
    result = insert_docx_at_content_control(main, "SB_DoesNotExist", sub_bytes)
    assert result is False


def test_insert_docx_placeholder_removed():
    """Original placeholder text should be gone after insert."""
    main = _make_doc_with_cc("SB_ObjectivesEndpoints")
    sub_bytes = _make_simple_docx("New content")
    insert_docx_at_content_control(main, "SB_ObjectivesEndpoints", sub_bytes)
    texts = [t.text for t in main.element.body.iter(f"{{{W_NS}}}t")]
    assert "placeholder" not in texts


def test_insert_docx_no_sectpr_copied():
    """The sectPr from the sub-document must not be inserted into the main doc body."""
    main = _make_doc_with_cc("SB_Flowchart")
    sub_bytes = _make_simple_docx("Flowchart content")
    insert_docx_at_content_control(main, "SB_Flowchart", sub_bytes)
    sect_prs = list(main.element.body.iter(f"{{{W_NS}}}sectPr"))
    # Main doc may have its own sectPr; the sub-doc's must not be duplicated
    assert len(sect_prs) <= 1


def test_insert_docx_multiple_paragraphs():
    """All paragraphs from the sub-document are inserted."""
    main = _make_doc_with_cc("SB_ObjectivesEndpoints")
    sub = Document()
    sub.add_paragraph("First paragraph")
    sub.add_paragraph("Second paragraph")
    sub.add_paragraph("Third paragraph")
    buf = io.BytesIO()
    sub.save(buf)

    result = insert_docx_at_content_control(main, "SB_ObjectivesEndpoints", buf.getvalue())
    assert result is True
    texts = [t.text for t in main.element.body.iter(f"{{{W_NS}}}t") if t.text]
    full = " ".join(texts)
    assert "First paragraph" in full
    assert "Second paragraph" in full
    assert "Third paragraph" in full


def test_insert_docx_content_control_preserved():
    """The <w:sdt> wrapper must remain around the inserted content."""
    main = _make_doc_with_cc("SB_Flowchart")
    sub_bytes = _make_simple_docx("Some content")
    insert_docx_at_content_control(main, "SB_Flowchart", sub_bytes)
    from elobs_word_updater.document.content_controls import find_content_control
    assert find_content_control(main, "SB_Flowchart") is not None


def test_insert_docx_paragraphs_wrap_table_content():
    """sdtContent must have an empty paragraph before AND after a leading/trailing
    table so Word renders both field markers outside the table rows."""
    main = _make_doc_with_cc("SB_SoA")
    sub = Document()
    sub.add_table(rows=1, cols=1)
    buf = io.BytesIO()
    sub.save(buf)

    insert_docx_at_content_control(main, "SB_SoA", buf.getvalue())

    from elobs_word_updater.document.content_controls import find_content_control
    sdt = find_content_control(main, "SB_SoA")
    sdt_content = sdt.find(f"{{{W_NS}}}sdtContent")
    children = list(sdt_content)
    # [p, tbl, p]
    assert children[0].tag == f"{{{W_NS}}}p",  "missing leading paragraph"
    assert children[1].tag == f"{{{W_NS}}}tbl", "table not second child"
    assert children[-1].tag == f"{{{W_NS}}}p",  "missing trailing paragraph"


def test_insert_docx_wrapping_paragraphs_always_added():
    """Leading and trailing empty paragraphs are added for any content type."""
    main = _make_doc_with_cc("SB_ObjectivesEndpoints")
    sub_bytes = _make_simple_docx("Just a paragraph")
    insert_docx_at_content_control(main, "SB_ObjectivesEndpoints", sub_bytes)

    from elobs_word_updater.document.content_controls import find_content_control
    sdt = find_content_control(main, "SB_ObjectivesEndpoints")
    sdt_content = sdt.find(f"{{{W_NS}}}sdtContent")
    children = list(sdt_content)
    assert children[0].tag == f"{{{W_NS}}}p",  "missing leading paragraph"
    assert children[-1].tag == f"{{{W_NS}}}p",  "missing trailing paragraph"


def test_study_000003_objectives_fixture_starts_with_table():
    """Validate the Study_000003 objectives fixture: body starts directly with a table
    (no leading paragraph), which is what requires the empty-paragraph fix."""
    fixture = FIXTURES / "objectives_endpoints_study_000003.docx"
    doc = Document(str(fixture))
    first = doc.element.body[0]
    assert first.tag == f"{{{W_NS}}}tbl", (
        "Study_000003 objectives fixture must start with a table — "
        "update the fixture if the API response has changed"
    )
    rows = list(first.iter(f"{{{W_NS}}}tr"))
    header_texts = [t.text or "" for t in rows[0].iter(f"{{{W_NS}}}t")]
    assert "Objectives" in header_texts
    assert "Endpoints" in header_texts


def test_objectives_study_000003_field_marker_outside_table():
    """When the Study_000003 objectives DOCX (which starts with a table) is inserted,
    the content control's first sdtContent child must be a paragraph so that Word
    renders the field marker before the table, not inside the first table row."""
    fixture = FIXTURES / "objectives_endpoints_study_000003.docx"
    sub_bytes = fixture.read_bytes()
    main = _make_doc_with_cc("SB_ObjectivesEndpoints")
    insert_docx_at_content_control(main, "SB_ObjectivesEndpoints", sub_bytes)

    from elobs_word_updater.document.content_controls import find_content_control
    sdt = find_content_control(main, "SB_ObjectivesEndpoints")
    sdt_content = sdt.find(f"{{{W_NS}}}sdtContent")
    children = list(sdt_content)

    assert children[0].tag == f"{{{W_NS}}}p", (
        "sdtContent must start with an empty paragraph so the field marker "
        "appears before the objectives table, not inside its first row"
    )
    assert children[1].tag == f"{{{W_NS}}}tbl", (
        "The objectives table must follow the empty paragraph inside sdtContent"
    )
    assert children[-1].tag == f"{{{W_NS}}}p", (
        "sdtContent must end with an empty paragraph so the field end marker "
        "appears after the objectives table, not inside its last row"
    )
