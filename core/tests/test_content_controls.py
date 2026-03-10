from __future__ import annotations
import io
import pytest
from lxml import etree
from docx import Document
from elobs_word_updater.document.content_controls import (
    find_content_control,
    set_content_control_text,
    set_content_control_lines,
    get_content_control_tags,
    clear_content_control,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _make_doc_with_cc(tag: str, initial_text: str = "placeholder") -> Document:
    """Create a minimal Document with one content control."""
    doc = Document()
    # Build the w:sdt XML and append to body
    xml = f"""
<w:sdt xmlns:w="{W_NS}">
  <w:sdtPr>
    <w:tag w:val="{tag}"/>
  </w:sdtPr>
  <w:sdtContent>
    <w:p><w:r><w:t>{initial_text}</w:t></w:r></w:p>
  </w:sdtContent>
</w:sdt>
""".strip()
    sdt = etree.fromstring(xml)
    doc.element.body.insert(0, sdt)
    return doc


def test_find_content_control_found():
    """Returns the sdt element when the tag is present in the document."""
    doc = _make_doc_with_cc("SB_ProtocolTitle")
    sdt = find_content_control(doc, "SB_ProtocolTitle")
    assert sdt is not None


def test_find_content_control_not_found():
    """Returns None when the tag does not exist in the document."""
    doc = _make_doc_with_cc("SB_ProtocolTitle")
    sdt = find_content_control(doc, "SB_DoesNotExist")
    assert sdt is None


def test_set_content_control_text():
    """Sets the text value of a content control and returns True."""
    doc = _make_doc_with_cc("SB_StudyID")
    result = set_content_control_text(doc, "SB_StudyID", "CDISC DEV-0")
    assert result is True
    sdt = find_content_control(doc, "SB_StudyID")
    texts = [t.text for t in sdt.iter(f"{{{W_NS}}}t")]
    assert "CDISC DEV-0" in texts


def test_set_content_control_text_missing_tag():
    """Returns False when the target tag is not in the document."""
    doc = _make_doc_with_cc("SB_StudyID")
    result = set_content_control_text(doc, "SB_NotPresent", "value")
    assert result is False


def test_set_content_control_lines():
    """Writes multiple lines as separate paragraphs and returns True."""
    doc = _make_doc_with_cc("SB_InclusionCriteria")
    lines = ["Criterion 1", "Criterion 2", "Criterion 3"]
    result = set_content_control_lines(doc, "SB_InclusionCriteria", lines)
    assert result is True
    sdt = find_content_control(doc, "SB_InclusionCriteria")
    texts = [t.text for t in sdt.iter(f"{{{W_NS}}}t")]
    assert texts == lines


def test_get_content_control_tags():
    """Returns all SB_* tags present in the document."""
    doc = Document()
    for tag in ["SB_ProtocolTitle", "SB_StudyID", "SB_Acronym"]:
        xml = f'<w:sdt xmlns:w="{W_NS}"><w:sdtPr><w:tag w:val="{tag}"/></w:sdtPr><w:sdtContent><w:p/></w:sdtContent></w:sdt>'
        doc.element.body.insert(0, etree.fromstring(xml))
    tags = get_content_control_tags(doc)
    assert set(tags) == {"SB_ProtocolTitle", "SB_StudyID", "SB_Acronym"}


def test_get_content_control_tags_excludes_non_sb():
    """Tags not prefixed with SB_ are not returned."""
    doc = Document()
    for tag in ["SB_ProtocolTitle", "MyCustomTag", "SB_StudyID"]:
        xml = f'<w:sdt xmlns:w="{W_NS}"><w:sdtPr><w:tag w:val="{tag}"/></w:sdtPr><w:sdtContent><w:p/></w:sdtContent></w:sdt>'
        doc.element.body.insert(0, etree.fromstring(xml))
    tags = get_content_control_tags(doc)
    assert "MyCustomTag" not in tags
    assert "SB_ProtocolTitle" in tags
    assert "SB_StudyID" in tags


def test_get_content_control_tags_empty_doc():
    """Returns an empty list when the document has no content controls."""
    doc = Document()
    assert get_content_control_tags(doc) == []


def test_find_content_control_no_sdtpr():
    """An sdt without sdtPr is silently skipped."""
    doc = Document()
    xml = f'<w:sdt xmlns:w="{W_NS}"><w:sdtContent><w:p/></w:sdtContent></w:sdt>'
    doc.element.body.insert(0, etree.fromstring(xml))
    assert find_content_control(doc, "SB_Anything") is None


def test_find_content_control_returns_first_match():
    """When two sdts share a tag, the first one in document order is returned."""
    doc = Document()
    for i in range(2):
        xml = f'<w:sdt xmlns:w="{W_NS}"><w:sdtPr><w:tag w:val="SB_ProtocolTitle"/></w:sdtPr><w:sdtContent><w:p><w:r><w:t>copy{i}</w:t></w:r></w:p></w:sdtContent></w:sdt>'
        doc.element.body.append(etree.fromstring(xml))
    sdt = find_content_control(doc, "SB_ProtocolTitle")
    texts = [t.text for t in sdt.iter(f"{{{W_NS}}}t")]
    assert "copy0" in texts


def test_set_content_control_text_empty_string():
    """Empty string is a valid value."""
    doc = _make_doc_with_cc("SB_StudyID")
    result = set_content_control_text(doc, "SB_StudyID", "")
    assert result is True
    sdt = find_content_control(doc, "SB_StudyID")
    texts = [t.text for t in sdt.iter(f"{{{W_NS}}}t")]
    assert "" in texts


def test_set_content_control_text_leading_trailing_space_sets_preserve():
    """Text with leading or trailing spaces gets xml:space='preserve'."""
    doc = _make_doc_with_cc("SB_StudyID")
    set_content_control_text(doc, "SB_StudyID", " spaced ")
    sdt = find_content_control(doc, "SB_StudyID")
    t_elems = list(sdt.iter(f"{{{W_NS}}}t"))
    assert any(
        e.get("{http://www.w3.org/XML/1998/namespace}space") == "preserve"
        for e in t_elems
    )


def test_set_content_control_text_replaces_previous_value():
    """Replaces existing text value without leaving the old text."""
    doc = _make_doc_with_cc("SB_StudyID", initial_text="old value")
    set_content_control_text(doc, "SB_StudyID", "new value")
    sdt = find_content_control(doc, "SB_StudyID")
    texts = [t.text for t in sdt.iter(f"{{{W_NS}}}t")]
    assert "old value" not in texts
    assert "new value" in texts


def test_set_content_control_lines_missing_tag():
    """Returns False when the target tag is not in the document."""
    doc = _make_doc_with_cc("SB_InclusionCriteria")
    result = set_content_control_lines(doc, "SB_DoesNotExist", ["line"])
    assert result is False


def test_set_content_control_lines_single_line():
    """Writes a single line and returns True."""
    doc = _make_doc_with_cc("SB_InclusionCriteria")
    result = set_content_control_lines(doc, "SB_InclusionCriteria", ["Only criterion"])
    assert result is True
    sdt = find_content_control(doc, "SB_InclusionCriteria")
    texts = [t.text for t in sdt.iter(f"{{{W_NS}}}t")]
    assert texts == ["Only criterion"]


def test_clear_content_control_removes_existing_content():
    """Removes all text runs from the content control and returns True."""
    doc = _make_doc_with_cc("SB_StudyID", initial_text="stale content from previous run")
    result = clear_content_control(doc, "SB_StudyID")
    assert result is True
    sdt = find_content_control(doc, "SB_StudyID")
    texts = [t.text for t in sdt.iter(f"{{{W_NS}}}t") if t.text]
    assert texts == []


def test_clear_content_control_leaves_empty_paragraph():
    """After clearing, sdtContent contains exactly one empty paragraph."""
    doc = _make_doc_with_cc("SB_StudyID", initial_text="old")
    clear_content_control(doc, "SB_StudyID")
    sdt = find_content_control(doc, "SB_StudyID")
    content = sdt.find(f"{{{W_NS}}}sdtContent")
    children = list(content)
    assert len(children) == 1
    assert children[0].tag == f"{{{W_NS}}}p"


def test_clear_content_control_missing_tag():
    """Returns False when the target tag is not in the document."""
    doc = _make_doc_with_cc("SB_StudyID")
    result = clear_content_control(doc, "SB_DoesNotExist")
    assert result is False


def test_set_content_control_lines_creates_one_paragraph_per_line():
    """Creates exactly one paragraph element per line."""
    doc = _make_doc_with_cc("SB_InclusionCriteria")
    lines = ["Line A", "Line B", "Line C"]
    set_content_control_lines(doc, "SB_InclusionCriteria", lines)
    sdt = find_content_control(doc, "SB_InclusionCriteria")
    paragraphs = list(sdt.iter(f"{{{W_NS}}}p"))
    assert len(paragraphs) == 3
