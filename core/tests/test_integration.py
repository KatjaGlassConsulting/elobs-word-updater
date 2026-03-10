from __future__ import annotations
import io
import pytest
import httpx
from pathlib import Path
from unittest.mock import AsyncMock, patch
from docx import Document
from lxml import etree
from elobs_word_updater.auth.no_auth import NoAuthClient
from elobs_word_updater.api.client import StudyApiClient
from elobs_word_updater.updater import update_document
from elobs_word_updater.document.content_controls import find_content_control
from tests.conftest import DUMMY_API_URL

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
FIXTURES = Path(__file__).parent / "fixtures"

ALL_TAGS = [
    "SB_ProtocolTitle", "SB_ProtocolTitleShort", "SB_StudyID", "SB_StudyPhase",
    "SB_EudraCTNumber", "SB_INDNumber", "SB_UniversalTrialNumber",
    "SB_InclusionCriteria", "SB_ExclusionCriteria",
    "SB_Flowchart", "SB_ObjectivesEndpoints", "SB_StudydesignGraphic",
]


def _make_template(tags: list[str]) -> str:
    """Create a minimal .docx template with the given SB_* content controls."""
    doc = Document()
    for tag in tags:
        xml = f"""
<w:sdt xmlns:w="{W_NS}">
  <w:sdtPr><w:tag w:val="{tag}"/></w:sdtPr>
  <w:sdtContent><w:p><w:r><w:t>PLACEHOLDER</w:t></w:r></w:p></w:sdtContent>
</w:sdt>""".strip()
        doc.element.body.append(etree.fromstring(xml))
    tmp_path = FIXTURES / "_test_template.docx"
    doc.save(str(tmp_path))
    return str(tmp_path)


@pytest.fixture
def api_client():
    auth = NoAuthClient(DUMMY_API_URL)
    return StudyApiClient(auth)


@pytest.fixture
def template_path():
    return _make_template(ALL_TAGS)


@pytest.fixture
def output_path(tmp_path):
    return str(tmp_path / "output.docx")


def _get_all_text(doc: Document) -> str:
    """Extract all text from a document including from sdt content."""
    texts = []
    for elem in doc.element.body.iter(f"{{{W_NS}}}t"):
        if elem.text:
            texts.append(elem.text)
    return " ".join(texts)


@pytest.mark.asyncio
async def test_update_latest_version(api_client, template_path, output_path):
    """Updating with version=None fetches the latest data and writes it to the document."""
    result = await update_document(
        doc_path=template_path,
        output_path=output_path,
        study_uid="Study_000001",
        version=None,
        api_client=api_client,
    )
    assert len(result.updated) > 0
    doc = Document(output_path)
    full_text = _get_all_text(doc)
    assert "Title for this Study (new)" in full_text


@pytest.mark.asyncio
async def test_update_version1(api_client, template_path, output_path):
    """Updating with version='1' fetches v1 data; title has no '(new)' suffix."""
    result = await update_document(
        doc_path=template_path,
        output_path=output_path,
        study_uid="Study_000001",
        version="1",
        api_client=api_client,
    )
    assert len(result.updated) > 0
    doc = Document(output_path)
    full_text = _get_all_text(doc)
    # v1 title has no "(new)" suffix
    assert "Title for this Study" in full_text
    assert "(new)" not in full_text


@pytest.mark.asyncio
async def test_specific_tags_only(api_client, template_path, output_path):
    """Only the requested tags are updated; other content controls remain untouched."""
    result = await update_document(
        doc_path=template_path,
        output_path=output_path,
        study_uid="Study_000001",
        version=None,
        api_client=api_client,
        tags=["SB_ProtocolTitle", "SB_StudyID"],
    )
    assert "SB_ProtocolTitle" in result.updated
    assert "SB_StudyID" in result.updated
    # Other tags should not be updated (still show as sdt in doc)
    doc = Document(output_path)
    assert find_content_control(doc, "SB_EudraCTNumber") is not None


@pytest.mark.asyncio
async def test_missing_tag_reported(api_client, output_path):
    """A tag requested but absent from the template is listed in result.missing."""
    # Template with only one tag
    template = _make_template(["SB_ProtocolTitle"])
    result = await update_document(
        doc_path=template,
        output_path=output_path,
        study_uid="Study_000001",
        version=None,
        api_client=api_client,
        tags=["SB_ProtocolTitle", "SB_StudyID"],
    )
    assert "SB_StudyID" in result.missing
    assert "SB_ProtocolTitle" in result.updated


@pytest.mark.asyncio
async def test_api_error_populates_api_errors(api_client, output_path):
    """When an API call fails, the label appears in api_errors and the tag is skipped."""
    template = _make_template(["SB_StudydesignGraphic"])
    with patch.object(
        api_client, "get_design_svg", new=AsyncMock(side_effect=httpx.ConnectError("timeout"))
    ):
        result = await update_document(
            doc_path=template,
            output_path=output_path,
            study_uid="Study_000001",
            version=None,
            api_client=api_client,
        )
    assert "design.svg" in result.api_errors
    assert "SB_StudydesignGraphic" in result.skipped


@pytest.mark.asyncio
async def test_empty_template_no_updates(api_client, output_path):
    """Template with no SB_ tags produces no updates and no errors."""
    template = _make_template([])
    result = await update_document(
        doc_path=template,
        output_path=output_path,
        study_uid="Study_000001",
        version=None,
        api_client=api_client,
    )
    assert len(result.updated) == 0
    assert len(result.missing) == 0
    assert len(result.api_errors) == 0


@pytest.mark.asyncio
async def test_inclusion_and_exclusion_criteria_populated(api_client, output_path):
    """Both inclusion and exclusion criteria tags are updated successfully."""
    template = _make_template(["SB_InclusionCriteria", "SB_ExclusionCriteria"])
    result = await update_document(
        doc_path=template,
        output_path=output_path,
        study_uid="Study_000001",
        version=None,
        api_client=api_client,
    )
    assert "SB_InclusionCriteria" in result.updated
    assert "SB_ExclusionCriteria" in result.updated


@pytest.mark.asyncio
async def test_sb_soa_tag_uses_flowchart(api_client, output_path):
    """SB_SoA tag is populated using the flowchart DOCX endpoint."""
    template = _make_template(["SB_SoA"])
    result = await update_document(
        doc_path=template,
        output_path=output_path,
        study_uid="Study_000001",
        version=None,
        api_client=api_client,
    )
    assert "SB_SoA" in result.updated


@pytest.mark.asyncio
async def test_all_text_tags_handled(api_client, output_path):
    """All text-based tags present in template are updated or skipped (never missing)."""
    text_tags = [
        "SB_ProtocolTitle", "SB_ProtocolTitleShort", "SB_StudyID", "SB_StudyPhase",
        "SB_EudraCTNumber", "SB_INDNumber", "SB_UniversalTrialNumber", "SB_Acronym",
    ]
    template = _make_template(text_tags)
    result = await update_document(
        doc_path=template,
        output_path=output_path,
        study_uid="Study_000001",
        version=None,
        api_client=api_client,
    )
    handled = set(result.updated) | set(result.skipped)
    for tag in text_tags:
        assert tag in handled, f"{tag} was neither updated nor skipped"


@pytest.mark.asyncio
async def test_output_file_is_valid_docx(api_client, template_path, output_path):
    """The output file can be opened as a valid DOCX."""
    await update_document(
        doc_path=template_path,
        output_path=output_path,
        study_uid="Study_000001",
        version=None,
        api_client=api_client,
    )
    doc = Document(output_path)
    assert doc is not None


@pytest.mark.asyncio
async def test_objectives_study_000003_field_outside_table(tmp_path):
    """Integration: objectives DOCX for Study_000003 starts with a table; the
    merger must place an empty paragraph before it so the field marker does not
    appear inside the table header row."""
    auth = NoAuthClient(DUMMY_API_URL)
    client = StudyApiClient(auth)
    template = _make_template(["SB_ObjectivesEndpoints"])
    out = str(tmp_path / "out.docx")

    result = await update_document(
        doc_path=template,
        output_path=out,
        study_uid="Study_000003",
        version=None,
        api_client=client,
        tags=["SB_ObjectivesEndpoints"],
    )

    assert "SB_ObjectivesEndpoints" in result.updated
    doc = Document(out)
    sdt = find_content_control(doc, "SB_ObjectivesEndpoints")
    sdt_content = sdt.find(f"{{{W_NS}}}sdtContent")
    children = list(sdt_content)
    assert children[0].tag == f"{{{W_NS}}}p",  "field marker anchor paragraph missing before table"
    assert children[1].tag == f"{{{W_NS}}}tbl", "objectives table not second child of sdtContent"
    assert children[-1].tag == f"{{{W_NS}}}p",  "trailing paragraph missing after objectives table"


@pytest.mark.asyncio
async def test_skipped_tag_content_is_cleared(tmp_path):
    """When the API returns no data for a tag, pre-existing content in the
    content control is cleared rather than left as stale output."""
    template = _make_template(["SB_StudydesignGraphic"])
    out = str(tmp_path / "out.docx")

    auth = NoAuthClient(DUMMY_API_URL)
    client = StudyApiClient(auth)

    with patch.object(client, "get_design_svg", new=AsyncMock(side_effect=httpx.ConnectError("timeout"))):
        result = await update_document(
            doc_path=template,
            output_path=out,
            study_uid="Study_000001",
            version=None,
            api_client=client,
        )

    assert "SB_StudydesignGraphic" in result.skipped
    doc = Document(out)
    sdt = find_content_control(doc, "SB_StudydesignGraphic")
    assert sdt is not None  # content control still present
    content = sdt.find(f"{{{W_NS}}}sdtContent")
    # Only the empty placeholder paragraph should remain — no stale text
    texts = [t.text for t in content.iter(f"{{{W_NS}}}t") if t.text]
    assert texts == []


@pytest.mark.asyncio
async def test_update_result_no_overlap(api_client, template_path, output_path):
    """A tag cannot appear in more than one of updated/skipped/missing."""
    result = await update_document(
        doc_path=template_path,
        output_path=output_path,
        study_uid="Study_000001",
        version=None,
        api_client=api_client,
    )
    all_tags = result.updated + result.skipped + result.missing
    assert len(all_tags) == len(set(all_tags)), "A tag appeared in multiple result buckets"
